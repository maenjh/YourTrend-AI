import os
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain_community.document_loaders import TextLoader
from langchain.prompts import PromptTemplate
from fpdf import FPDF, HTMLMixin
import markdown
from docx import Document

# .env 파일에서 환경 변수 로드
load_dotenv()

# OpenAI API 키를 환경 변수에서 가져옴
openai_api_key = os.getenv("OPENAI_API_KEY")

# API 키가 없으면 오류 발생
if not openai_api_key:
    raise ValueError("OPENAI_API_KEY가 설정되지 않았습니다. .env 파일을 확인해주세요.")

# ChatGPT-4o-Mini 모델 초기화
llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.7, openai_api_key=openai_api_key)

def generate_report(project_name, team_members, topic, template=None):
    """
    주어진 프로젝트 정보를 바탕으로 보고서를 생성하는 함수
    """
    if template:
        # 템플릿의 각 섹션을 채우도록 요청
        sections = template.split('\n\n')  # 빈 줄을 기준으로 섹션 구분
        filled_sections = []
        for section in sections:
            prompt = PromptTemplate(
                input_variables=["project_name", "team_members", "topic", "section"],
                template="""
                프로젝트명: {project_name}
                팀원: {team_members}
                주제: {topic}
                
                다음 섹션을 채워주세요:
                
                {section}
                
                위 정보를 바탕으로 이 섹션에 대한 구체적이고 전문적인 내용을 작성해주세요.
                """
            )
            query = prompt.format(project_name=project_name, team_members=team_members, topic=topic, section=section)
            chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=Chroma(embedding_function=OpenAIEmbeddings()).as_retriever(),
                return_source_documents=True
            )
            result = chain.invoke({"query": query})
            filled_sections.append(result["result"])
        
        return "\n\n".join(filled_sections)
    else:
        # 기존의 기본 템플릿 사용
        prompt = PromptTemplate(
            input_variables=["project_name", "team_members", "topic"],
            template="""
            프로젝트명: {project_name}
            팀원: {team_members}
            주제: {topic}

            위 정보를 바탕으로 다음 구조를 가진 상세한 보고서를 작성해주세요:

            1. 개요
            2. 프로젝트 목표
            3. 팀원 소개 및 역할
            4. 방법론
            5. 주요 발견사항
            6. 결론 및 향후 계획

            각 섹션에 대해 구체적이고 전문적인 내용을 포함시켜 주세요.
            """
        )
        query = prompt.format(project_name=project_name, team_members=team_members, topic=topic)
        chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=Chroma(embedding_function=OpenAIEmbeddings()).as_retriever(),
            return_source_documents=True
        )
        result = chain.invoke({"query": query})
        return result["result"]

def save_report(content, filename, format):
    """
    생성된 보고서를 지정된 형식으로 저장하는 함수
    """
    os.makedirs("reports", exist_ok=True)
    filepath = os.path.join("reports", f"{filename}.{format}")

    if format == "md":
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
    elif format == "pdf":
        class PDF(FPDF, HTMLMixin):
            pass
        html = markdown.markdown(content)
        pdf = PDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.write_html(html)
        pdf.output(filepath)
    elif format == "docs":  # DOCS 형식 지원
        doc = Document()
        doc.add_heading(filename.replace('_', ' ').title(), level=1)
        doc.add_paragraph(content)
        doc.save(filepath)
    else:
        raise ValueError("지원되지 않는 형식입니다. 'md', 'pdf', 또는 'docs'를 사용하세요.")

    print(f"보고서가 {filepath}로 저장되었습니다.")

def load_template(template_path):
    """
    지정된 경로에서 보고서 템플릿을 로드하는 함수
    """
    if not os.path.exists(template_path):
        print(f"템플릿 파일을 찾을 수 없습니다: {template_path}")
        return None

    file_extension = os.path.splitext(template_path)[1].lower()

    try:
        if file_extension == '.docx':
            doc = Document(template_path)
            sections = []
            current_section = []
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    if current_section:
                        sections.append('\n'.join(current_section))
                        current_section = []
                current_section.append(paragraph.text)
            if current_section:
                sections.append('\n'.join(current_section))
            return '\n\n'.join(sections)
        elif file_extension in ['.txt', '.md']:
            with open(template_path, 'r', encoding='utf-8') as file:
                return file.read()
        else:
            print(f"지원되지 않는 파일 형식입니다: {file_extension}")
            return None
    except Exception as e:
        print(f"파일을 읽는 중 오류가 발생했습니다: {str(e)}")
        return None

def main():
    """
    메인 함수: 사용자 입력을 받아 보고서를 생성하고 저장
    """
    # 보고서 양식 선택
    while True:
        choice = input("보고서 양식을 선택하세요 (1: 기본 양식, 2: 사용자 정의 양식): ")
        if choice in ['1', '2']:
            break
        print("잘못된 선택입니다. 1 또는 2를 입력하세요.")
    
    template = None
    if choice == '2':
        template_path = input("사용자 정의 양식 파일의 경로를 입력하세요: ")
        template = load_template(template_path)
        if not template:
            print("양식 로드에 실패했습니다. 기본 양식을 사용합니다.")
    
    # 보고서 형식 선택
    while True:
        format = input("보고서 형식을 선택하세요 (md/pdf/docs): ").lower()
        if format in ["md", "pdf", "docs"]:
            break
        print("잘못된 형식입니다. 다시 입력하세요: md, pdf 또는 docs를 선택하세요.")
    
    # 프로젝트 정보 입력
    project_name = input("프로젝트명을 입력하세요: ")
    team_members = input("팀원들의 이름을 쉼표로 구분하여 입력하세요: ")
    topic = input("프로젝트 주제를 입력하세요: ")
    
    # 보고서 생성
    report_content = generate_report(project_name, team_members, topic, template)
    
    # 보고서 저장
    filename = f"{project_name.replace(' ', '_')}_report"
    save_report(report_content, filename, format)

if __name__ == "__main__":
    main()